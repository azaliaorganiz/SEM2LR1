import docx
from time import time
from fuzzywuzzy import fuzz
import Levenshtein
from docx import Document


file1 = docx.Document("text1.docx")
file2 = docx.Document("text2.docx")


text1 = "\n".join([paragraph.text for paragraph in file1.paragraphs])
text2 = "\n".join([paragraph.text for paragraph in file2.paragraphs])

time1 = time()
distance = Levenshtein.distance(text1, text2)
time2 = time()

#расчет коэффициента сходства (от 0 до 1) и перевод в проценты
shodstvo = Levenshtein.ratio(text1, text2)*100

print(f"Расстояние Левенштейна: {distance}")
print(f"Схожесть (%): {shodstvo} %")
print(f"Вычислено за: {round(time2 - time1, 4)} сек")


time3 = time()
ratio = fuzz.ratio(text1, text2)
time4 = time()
print(f"Расчёт через FuzzyWuzzy: {ratio} % ")
print(f"Вычислено за: {round(time4 - time3, 4)} сек.")


def levenshtein(text1, text2):
    n, m = len(text1), len(text2)
    if n > m:
        text1, text2 = text2, text1
        n, m = m, n

    current_row = range(n + 1)
    for i in range(1, m + 1):
        previous_row, current_row = current_row, [i] + [0] * n
        for j in range(1, n + 1):
            add, delete, change = previous_row[j] + 1, current_row[j - 1] + 1, previous_row[j - 1]
            if text1[j-1] != text2[i-1]:
                change += 1
            current_row[j] = min(add, delete, change)
    return current_row[n]

time5 = time()
custom_distance = levenshtein(text1, text2)
time6 = time()

shodstvo = Levenshtein.ratio(text1, text2)*100
print(f"Самостоятельный расчет расстояния Левенштейна(через динамическое програмирование): {custom_distance}")
print(f"Схожесть (%): {shodstvo} %")
print(f"Вычислено за: {round(time6 - time5, 4)} сек.")

